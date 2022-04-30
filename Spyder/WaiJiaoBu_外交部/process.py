"""
利用docx将文本转为docx
"""

import json
import docx

contents = []
with open("./wjb_2020.txt", "r", encoding='utf-8') as f:
    for line in f.readlines():
        contents.append(json.loads(line))

# 过滤 2020 年的记录，按照时间顺序排列
contents = list(filter(lambda x: '2020' in x['date'], contents))
contents = sorted(contents, key=lambda x: x["date"])

doc = docx.Document()

for c in contents:
    date, content = c["date"], c["text"]
    try:
        title = content[:26].strip()
        doc.add_paragraph(title+"\n\n")
        body = content.split("打印本页")[1].split("相关新闻")[0]
        run = doc.add_paragraph(body)
        doc.add_page_break()        # 分页
    except Exception as e:
        print(e)

doc.save('wjb.docx')

